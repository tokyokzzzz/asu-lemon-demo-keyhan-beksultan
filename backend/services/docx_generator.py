import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── XML helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_auto_bg(cell):
    """Set cell background to 'auto' — inherits document/theme background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd element first
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'auto')
    tcPr.append(shd)


def _set_auto_color(run):
    """Set font color to 'auto' — Word picks dark/light automatically based on background."""
    rPr = run._r.get_or_add_rPr()
    # Remove any existing color element
    for existing in rPr.findall(qn('w:color')):
        rPr.remove(existing)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), 'auto')
    rPr.append(color_el)


def _add_run(para, text: str, bold=False, font_size=11,
             color_hex=None, font_name='Times New Roman'):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color_hex:
        r = int(color_hex[0:2], 16)
        g = int(color_hex[2:4], 16)
        b = int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    else:
        _set_auto_color(run)
    return run


# ─── Parser ───────────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Remove markdown table artifacts (| symbols and separator lines) from AI output."""
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        # Skip pure separator lines like |---|---| or |-----------|
        if re.match(r'^[\s|:\-]+$', line):
            continue
        # Strip leading/trailing | and spaces, then clean inner | used as column separators
        line = line.strip().strip('|').strip()
        # Replace remaining | that act as column separators with a space
        line = line.replace('|', ' ').strip()
        # Collapse multiple spaces
        line = re.sub(r'  +', ' ', line)
        cleaned.append(line)
    return '\n'.join(cleaned)


_SECTION_RE = re.compile(
    r'^(РАЗДЕЛ|Раздел|БӨЛІМ|Бөлім|Section)\s+\d',
    re.IGNORECASE
)
_SUBSECTION_RE = re.compile(r'^(\d+\.\d+(?:\.\d+)?)[.\s:]\s*(.*)')


def _parse_tz_sections(text: str):
    """
    Parse TZ text into a list of sections.

    Returns:
        [
          {
            'title': 'РАЗДЕЛ 1: ОБЩИЕ СВЕДЕНИЯ',
            'subsections': [
              {'num': '1.1', 'name': 'Наименование приоритета', 'content': '...'},
              {'num': '1.2', 'name': '...', 'content': '...'},
            ]
          },
          ...
        ]
    """
    lines = text.split('\n')
    sections = []
    current_section = None
    current_sub = None

    def flush_sub():
        if current_sub and current_section is not None:
            current_section['subsections'].append({
                'num': current_sub['num'],
                'name': current_sub['name'],
                'content': '\n'.join(current_sub['lines']).strip()
            })

    for line in lines:
        s = line.strip()
        if not s:
            if current_sub:
                current_sub['lines'].append('')
            continue

        if _SECTION_RE.match(s):
            flush_sub()
            current_sub = None
            current_section = {'title': s, 'subsections': []}
            sections.append(current_section)

        elif m := _SUBSECTION_RE.match(s):
            flush_sub()
            current_sub = {
                'num': m.group(1),
                'name': m.group(2).strip(),
                'lines': []
            }

        else:
            if current_sub:
                current_sub['lines'].append(s)
            elif current_section:
                # Content before any subsection — attach to section title area
                if not current_section['subsections']:
                    current_section.setdefault('intro', '')
                    current_section['intro'] = (current_section.get('intro', '') + '\n' + s).strip()
            else:
                # Text before first section header — create implicit section
                current_section = {'title': s, 'subsections': []}
                sections.append(current_section)

    flush_sub()
    return sections


# ─── Table builder ────────────────────────────────────────────────────────────

# No explicit color — all text uses 'auto' color (dark on light bg, white on dark bg)


def _build_tz_table(doc: Document, sections: list):
    """
    Build a single-column table: one row per section.
    Each cell contains the section title + all its subsections.
    """
    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'

    # Fix table width to full page width
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9638')   # ~17cm in twentieths of a point
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    for section in sections:
        row = table.add_row()
        cell = row.cells[0]
        cell.text = ''   # clear default paragraph
        _set_cell_auto_bg(cell)  # transparent — inherits document theme background

        # ── Section title paragraph ──
        title_para = cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(title_para, section['title'], bold=True, font_size=11)

        # Optional intro text (content before subsections)
        if section.get('intro'):
            p = cell.add_paragraph()
            _add_run(p, section['intro'], font_size=11)

        # ── Subsection paragraphs ──
        for sub in section['subsections']:
            # Subsection header line: "1.1  Наименование показателя"
            sub_header = cell.add_paragraph()
            _add_run(sub_header, f"{sub['num']}  {sub['name']}",
                     bold=True, font_size=11)

            # Subsection content (may be multi-line)
            if sub['content']:
                for line in sub['content'].split('\n'):
                    line = line.strip()
                    if line:
                        content_para = cell.add_paragraph()
                        content_para.paragraph_format.left_indent = Cm(0.5)
                        _add_run(content_para, line, font_size=11)

        # Cell padding
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ('top', 'left', 'bottom', 'right'):
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), '80')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    return table


# ─── Public API ───────────────────────────────────────────────────────────────

def create_docx_from_text(text: str, filename: str, language: str) -> str:
    """
    Create a formatted DOCX with a single-column table.
    Rows = sections (1 row per section, subsections inside the cell).
    """
    doc = Document()

    # A4 page margins
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(1.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    set_document_language(doc, 'kk' if language == 'kazakh' else 'ru')

    # Document title — outside the table, at the very top
    title_text = 'ТЕХНИЧЕСКОЕ ЗАДАНИЕ' if language != 'kazakh' else 'ТЕХНИКАЛЫҚ ТАПСЫРМА'
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Cm(0)
    title_para.paragraph_format.space_after = Cm(0.6)
    run = _add_run(title_para, title_text, bold=True, font_size=16)
    run.underline = True  # visual separator from table

    # Parse and build
    sections = _parse_tz_sections(_clean_text(text))
    if sections:
        _build_tz_table(doc, sections)
    else:
        doc.add_paragraph(text)

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = filename.rsplit('.', 1)[0]
    output_path = f"uploads/document_{timestamp}_{base_name}.docx"
    doc.save(output_path)
    return output_path


def set_document_language(doc: Document, language_code: str) -> None:
    try:
        doc.core_properties.language = language_code
    except Exception:
        pass


# ── Legacy ────────────────────────────────────────────────────────────────────

def generate_corrected_docx(original_file_path: str, corrected_sections: dict, language: str) -> str:
    original_filename = Path(original_file_path).name
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f"uploads/corrected_{timestamp}_{original_filename.rsplit('.', 1)[0]}.docx"
    combined = '\n\n'.join(f"{k}\n{v}" for k, v in corrected_sections.items())
    doc = Document()
    set_document_language(doc, 'kk' if language == 'kazakh' else 'ru')
    sections = _parse_tz_sections(combined)
    _build_tz_table(doc, sections)
    doc.save(output_path)
    return output_path


def _create_docx_from_corrected_sections(corrected_sections: dict, language: str) -> Document:
    combined = '\n\n'.join(f"{k}\n{v}" for k, v in corrected_sections.items())
    doc = Document()
    set_document_language(doc, 'kk' if language == 'kazakh' else 'ru')
    sections = _parse_tz_sections(combined)
    _build_tz_table(doc, sections)
    return doc
