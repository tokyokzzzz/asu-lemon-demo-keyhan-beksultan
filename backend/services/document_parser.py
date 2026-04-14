import re
from pathlib import Path


def extract_text_from_file(file_path: str, filename: str) -> str:
    """
    Extract text from DOCX or PDF files.

    Args:
        file_path: Full path to the file
        filename: Original filename to determine format

    Returns:
        Extracted text with normalized whitespace

    Raises:
        ValueError: If file format is not supported
    """
    file_extension = filename.lower().split('.')[-1]

    if file_extension == 'docx':
        return _extract_from_docx(file_path)
    elif file_extension == 'pdf':
        return _extract_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .docx, .pdf")


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            paragraph = [p for p in doc.paragraphs if p._element == element][0]
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        elif element.tag.endswith('tbl'):
            table = [t for t in doc.tables if t._element == element][0]
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

    text = "\n".join(text_parts)
    return _normalize_whitespace(text)


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using pdfplumber, fallback to PyMuPDF."""
    import pdfplumber
    import fitz

    text = ""

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception:
        pass

    if len(text.strip()) < 100:
        try:
            pdf_document = fitz.open(file_path)
            text = ""
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n"
            pdf_document.close()
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    return _normalize_whitespace(text)


def _normalize_whitespace(text: str) -> str:
    """Normalize excessive whitespace in text."""
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    text = text.strip()
    return text


def detect_document_structure(text: str) -> dict:
    """
    Detect which required sections are present in the document.

    Args:
        text: Document text to analyze

    Returns:
        Dictionary with boolean flags for each section
    """
    text_lower = text.lower()

    structure = {
        "has_section_1": _contains_any(text_lower, [
            "общие сведения",
            "жалпы мәліметтер"
        ]),
        "has_section_2": _contains_any(text_lower, [
            "цели и задачи",
            "бағдарламаның мақсаттары"
        ]),
        "has_section_3": _contains_any(text_lower, [
            "стратегических и программных документов",
            "стратегиялық"
        ]),
        "has_section_4_1": _contains_any(text_lower, [
            "прямые результаты",
            "тікелей нәтижелер"
        ]),
        "has_section_4_2": _contains_any(text_lower, [
            "конечный результат",
            "түпкі нәтиже"
        ]),
        "has_section_5": _contains_any(text_lower, [
            "предельная сумма",
            "шекті сома"
        ])
    }

    return structure


def _contains_any(text: str, keywords: list) -> bool:
    """Check if text contains any of the keywords."""
    for keyword in keywords:
        if keyword in text:
            return True
    return False
